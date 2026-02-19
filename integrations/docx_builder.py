from docx import Document


def build_pr_docx(
    *,
    header: str,
    title: str,
    title_lines: list[str],
    body_lines: list[str],
    output_path: str,
) -> None:
    doc = Document()

    if (header or "").strip():
        p = doc.add_paragraph()
        run = p.add_run(header.strip())
        run.bold = True

    normalized_title_lines = [(x or "").strip() for x in (title_lines or []) if (x or "").strip()]
    if not normalized_title_lines:
        fallback_title = (title or "").strip() or "新聞稿"
        normalized_title_lines = [fallback_title]

    for tl in normalized_title_lines:
        p = doc.add_paragraph(tl)
        try:
            p.style = "Heading 1"
        except Exception:
            run = p.runs[0] if p.runs else p.add_run(tl)
            run.bold = True

    # 标题与正文之间留一空行，符合下游格式要求
    doc.add_paragraph("")

    for line in body_lines or []:
        doc.add_paragraph((line or "").rstrip())

    doc.save(output_path)
